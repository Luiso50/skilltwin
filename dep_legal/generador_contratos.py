import os
import logging
from datetime import datetime
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from dep_legal.header_documento import crear_documento_con_encabezado

logger = logging.getLogger('skilltwin.contratos')

CONTRATOS_DIR = os.path.join(os.path.dirname(__file__), "contratos")


def generar_contrato(id_experto, nombre, especialidad, comision=15.0):
    """Genera un acuerdo de servicio legal en formato Word (.docx) para un nuevo experto."""
    # Asegurar que la carpeta de contratos existe
    if not os.path.exists(CONTRATOS_DIR):
        os.makedirs(CONTRATOS_DIR)

    nombre_archivo = f"contrato_{id_experto}.docx"
    ruta_guardado = os.path.join(CONTRATOS_DIR, nombre_archivo)

    try:
        # Crear documento con encabezado estándar SkillTwin
        doc = crear_documento_con_encabezado()

        # ── TÍTULO DEL CONTRATO ──
        titulo = doc.add_heading(
            "ACUERDO DE LICENCIA DE CLON DIGITAL Y SERVICIOS", level=0
        )
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run("SKILLTWIN")
        run_sub.font.size = Pt(14)
        run_sub.font.color.rgb = RGBColor(79, 123, 186)
        run_sub.bold = True

        doc.add_paragraph()  # Espaciado

        # ── FECHA ──
        p_fecha = doc.add_paragraph()
        run_fecha = p_fecha.add_run(
            f"Con fecha de hoy, {datetime.now().strftime('%d/%m/%Y')}, las partes acuerdan:"
        )
        run_fecha.italic = True

        doc.add_paragraph()

        # ── SECCIÓN 1: PARTES ──
        doc.add_heading("1. PARTES CONTRATANTES", level=2)

        p_parte1 = doc.add_paragraph()
        p_parte1.add_run("De una parte: ").bold = True
        p_parte1.add_run(
            "La plataforma SKILLTWIN (en adelante, la \"Plataforma\")."
        )

        p_parte2 = doc.add_paragraph()
        p_parte2.add_run("De otra parte: ").bold = True
        p_parte2.add_run(
            f"D/Dña. {nombre}, especialista en {especialidad} "
            f"(en adelante, el \"Licenciante\")."
        )

        doc.add_paragraph()

        # ── SECCIÓN 2: OBJETO ──
        doc.add_heading("2. OBJETO DEL ACUERDO", level=2)

        doc.add_paragraph(
            "El Licenciante concede una licencia no exclusiva y revocable a la Plataforma "
            "para procesar su base de conocimiento provista y generar un \"Gemelo Digital\" "
            "(Clon de IA) capaz de responder preguntas en su nombre a usuarios de la red."
        )

        doc.add_paragraph()

        # ── SECCIÓN 3: COMISIONES ──
        doc.add_heading("3. COMISIONES Y FACTURACIÓN", level=2)

        comision_texto = [
            "- La Plataforma cobrará una tarifa a los clientes finales por cada consulta "
            "realizada al Clon de IA.",
            f"- De los ingresos generados, la Plataforma retendrá un {comision}% en concepto "
            f"de comisión por servicio, mantenimiento de servidores y procesamiento de APIs.",
            f"- El {100 - comision}% restante será transferido al Licenciante de forma mensual.",
        ]

        for linea in comision_texto:
            doc.add_paragraph(linea)

        doc.add_paragraph()

        # ── SECCIÓN 4: PROTECCIÓN DE DATOS ──
        doc.add_heading("4. PROTECCIÓN DE DATOS Y PRIVACIDAD", level=2)

        privacidad = [
            "- La Plataforma se compromete a no compartir, transferir ni utilizar la base "
            "de conocimiento del Licenciante para entrenar otros modelos de IA externos.",
            "- El Licenciante puede solicitar la baja total del servicio y la eliminación "
            "completa de sus datos y de su Clon de IA en cualquier momento con un "
            "preaviso de 48 horas.",
        ]

        for linea in privacidad:
            doc.add_paragraph(linea)

        doc.add_paragraph()

        # ── SECCIÓN 5: RESPONSABILIDAD ──
        doc.add_heading("5. RESPONSABILIDAD", level=2)

        doc.add_paragraph(
            "La Plataforma no será responsable de las opiniones o respuestas generadas "
            "por el Clon de IA, las cuales tienen carácter estrictamente consultivo e informal."
        )

        doc.add_paragraph()
        doc.add_paragraph()

        # ── FIRMAS ──
        p_firma_titulo = doc.add_paragraph()
        p_firma_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ft = p_firma_titulo.add_run("FIRMADO DIGITALMENTE EN CONFORMIDAD:")
        run_ft.bold = True
        run_ft.font.size = Pt(12)

        doc.add_paragraph()

        # Tabla de firmas
        tabla_firmas = doc.add_table(rows=2, cols=2)
        tabla_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # SkillTwin
        cell_st = tabla_firmas.cell(0, 0)
        p_st = cell_st.paragraphs[0]
        p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_st = p_st.add_run("SKILLTWIN CORP")
        run_st.bold = True
        run_st.font.size = Pt(10)

        p_st2 = cell_st.add_paragraph()
        p_st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_st2.add_run("Representante de la Plataforma").font.size = Pt(9)

        # Licenciante
        cell_lic = tabla_firmas.cell(0, 1)
        p_lic = cell_lic.paragraphs[0]
        p_lic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_lic = p_lic.add_run(nombre.upper())
        run_lic.bold = True
        run_lic.font.size = Pt(10)

        p_lic2 = cell_lic.add_paragraph()
        p_lic2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lic2.add_run("Licenciante del Clon").font.size = Pt(9)

        # ── PIE DE PÁGINA ──
        doc.add_paragraph()
        doc.add_paragraph()

        p_pie = doc.add_paragraph()
        p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pie = p_pie.add_run(
            "SkillTwin © 2026 — Gemelos Digitales Profesionales"
        )
        run_pie.font.size = Pt(8)
        run_pie.font.color.rgb = RGBColor(150, 150, 150)
        run_pie.italic = True

        # Guardar documento
        doc.save(ruta_guardado)

        logger.info(f"Contrato de licencia generado: {ruta_guardado}")
        return ruta_guardado

    except Exception as e:
        logger.error(f"No se pudo generar el contrato: {e}")
        return None


def main():
    logger.info("GENERADOR DE CONTRATOS LEGALES - SKILLTWIN")

    nombre = input("Nombre completo del profesional: ").strip()
    id_experto = input("ID de usuario único (ej. rsanchez): ").strip().lower()
    especialidad = input("Especialidad o Habilidad a licenciar: ").strip()

    comision_str = input(
        "Porcentaje de comisión para la plataforma (Defecto: 15%): "
    ).strip()
    comision = 15.0
    if comision_str:
        try:
            comision = float(comision_str)
        except ValueError:
            logger.warning("Entrada inválida, usando 15.0% por defecto")

    if nombre and id_experto and especialidad:
        generar_contrato(id_experto, nombre, especialidad, comision)
    else:
        logger.warning("Todos los campos de texto son obligatorios")


if __name__ == "__main__":
    main()
