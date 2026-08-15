"""
Módulo para agregar encabezados profesionales a documentos Word.
Todos los documentos SkillTwin usarán este módulo para el encabezado estándar.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Rutas de imágenes
LOGO_DIR = os.path.join(os.path.dirname(__file__), "..", "cerebro")
LOGO_PATH = os.path.join(LOGO_DIR, "logo-mark.svg")
LOGO_PNG_PATH = os.path.join(LOGO_DIR, "logo-mark.png")

# Información de la empresa
EMPRESA_INFO = {
    "nombre": "SKILLTWIN",
    "direccion": "Calle Gran Vía, 28, 4ª Izq, 28013 Madrid, España",
    "telefono": "+34 600 000 000",
    "web": "https://skilltwin.es",
    "email": "teamskiltwinhq@zohomail.com"
}


def agregar_encabezado(document: Document):
    """
    Agrega un encabezado profesional con logo a la izquierda
    e información de contacto a la derecha.
    """
    header = document.sections[0].header
    header.is_linked_to_previous = False

    # Crear tabla de 1 fila x 2 columnas para el encabezado
    table = header.add_table(rows=1, cols=2, width=Cm(16))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Ancho de columnas
    table.columns[0].width = Cm(6)   # Logo + Nombre
    table.columns[1].width = Cm(10)  # Info contacto

    # ── CELDA IZQUIERDA: Logo + Nombre ──
    cell_izq = table.cell(0, 0)
    cell_izq.paragraphs[0].clear()

    # Buscar imagen del logo (solo PNG funciona en Word)
    logo_insertado = False
    if os.path.exists(LOGO_PNG_PATH):
        try:
            p_logo = cell_izq.paragraphs[0]
            run_logo = p_logo.add_run()
            run_logo.add_picture(LOGO_PNG_PATH, width=Cm(1.5))
            logo_insertado = True
        except Exception as e:
            print(f"[HEADER] Error al insertar logo: {e}")

    # Si no hay imagen, usar representación de texto del logo
    if not logo_insertado:
        p_logo = cell_izq.paragraphs[0]
        run_logo = p_logo.add_run("ST")
        run_logo.bold = True
        run_logo.font.size = Pt(16)
        run_logo.font.color.rgb = RGBColor(79, 123, 186)  # Color azul de SkillTwin

    # Nombre de la empresa debajo del logo
    p_nombre = cell_izq.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_nombre = p_nombre.add_run(EMPRESA_INFO["nombre"].upper())
    run_nombre.bold = True
    run_nombre.font.size = Pt(11)
    run_nombre.font.color.rgb = RGBColor(33, 37, 41)

    # ── CELDA DERECHA: Información de contacto ──
    cell_der = table.cell(0, 1)
    cell_der.paragraphs[0].clear()

    # Dirección
    p_dir = cell_der.paragraphs[0]
    p_dir.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_dir = p_dir.add_run(EMPRESA_INFO['direccion'])
    run_dir.font.size = Pt(8)
    run_dir.font.color.rgb = RGBColor(100, 100, 100)

    # Teléfono y Email
    p_tel = cell_der.add_paragraph()
    p_tel.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_tel = p_tel.add_run(f"Tel: {EMPRESA_INFO['telefono']}  |  Email: {EMPRESA_INFO['email']}")
    run_tel.font.size = Pt(8)
    run_tel.font.color.rgb = RGBColor(100, 100, 100)

    # Web
    p_web = cell_der.add_paragraph()
    p_web.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_web = p_web.add_run(EMPRESA_INFO['web'])
    run_web.font.size = Pt(8)
    run_web.font.color.rgb = RGBColor(79, 123, 186)  # Color azul SkillTwin

    # Separador debajo del encabezado
    p_sep = header.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = p_sep.add_run("─" * 80)
    run_sep.font.size = Pt(6)
    run_sep.font.color.rgb = RGBColor(200, 200, 200)

    return document


def crear_documento_con_encabezado() -> Document:
    """
    Crea un nuevo documento Word con el encabezado estándar de SkillTwin.
    """
    doc = Document()

    # Configurar estilos base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Agregar encabezado
    agregar_encabezado(doc)

    return doc


if __name__ == "__main__":
    # Prueba rápida: generar documento de ejemplo
    doc = crear_documento_con_encabezado()

    doc.add_heading("Documento de Prueba", level=1)
    doc.add_paragraph("Este es un documento de prueba con encabezado SkillTwin.")

    test_path = os.path.join(os.path.dirname(__file__), "test_documento.docx")
    doc.save(test_path)
    print(f"Documento de prueba guardado en: {test_path}")
