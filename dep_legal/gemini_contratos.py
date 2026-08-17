"""
Módulo para generar contratos automáticos en formato Word (.docx).
Se integra con el orquestador de órdenes y usa el encabezado estándar SkillTwin.
"""

import os
import json
import urllib.request
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from dep_legal.header_documento import crear_documento_con_encabezado


def guardar_contrato_docx(doc: Document, orden_id: str, clon_id: str) -> str:
    """
    Guarda un documento Word en la carpeta de contratos.
    Retorna la ruta del archivo guardado.
    """
    contratos_dir = os.path.join(os.path.dirname(__file__), "contratos")
    if not os.path.exists(contratos_dir):
        os.makedirs(contratos_dir)

    nombre_archivo = f"contrato_{orden_id}_{clon_id}.docx"
    ruta_guardado = os.path.join(contratos_dir, nombre_archivo)

    doc.save(ruta_guardado)
    return ruta_guardado


def generar_contrato_gemini(
    orden_id,
    cliente_email,
    clon_id,
    clon_nombre,
    especialidad,
    cantidad_horas,
    monto_total,
    comision,
):
    """
    Genera un contrato legal profesional en formato Word (.docx).
    Usa Gemini AI para el contenido y el encabezado estándar SkillTwin.
    Retorna la ruta del archivo guardado.
    """

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return generar_contrato_default(
            orden_id,
            cliente_email,
            clon_id,
            clon_nombre,
            especialidad,
            cantidad_horas,
            monto_total,
            comision,
        )

    model_name = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent"

    fecha_inicio = datetime.now()
    fecha_fin = fecha_inicio + timedelta(days=30)

    prompt = f"""
Eres un abogado especializado en contratos de licenciamiento de propiedad intelectual.
Genera un contrato profesional y legal para los siguientes datos:

DATOS DEL CONTRATO:
- ID de Orden: {orden_id}
- Fecha: {fecha_inicio.strftime("%d de %B de %Y")}
- Cliente: {cliente_email}
- Experto (Licenciante): {clon_nombre}
- Especialidad: {especialidad}
- Horas de Servicio: {cantidad_horas} horas
- Tarifa: ${monto_total / cantidad_horas:.2f}/hora
- Total a Pagar: ${monto_total:.2f}
- Comisión de Plataforma: ${comision:.2f}
- Fecha de Vencimiento: {fecha_fin.strftime("%d de %B de %Y")}

Genera un contrato en español que incluya:
1. Título del contrato
2. Partes del Contrato (Cliente y Licenciante)
3. Descripción de Servicios (Artículo 1)
4. Términos Económicos (Artículo 2)
5. Plazo (Artículo 3)
6. Confidencialidad (Artículo 4)
7. Limitación de Responsabilidad (Artículo 5)
8. Terminación (Artículo 6)
9. Ley Aplicable (Artículo 7)
10. Firmas

Responde SOLO con el contenido del contrato, sin explicaciones adicionales.
Incluye números de artículos para referencia.
"""

    headers = {"Content-Type": "application/json", "x-goog-api-key": api_key}
    body = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.3,
            "maxOutputTokens": 2000,
            "responseMimeType": "text/plain",
        },
    }

    try:
        req = urllib.request.Request(
            url, data=json.dumps(body).encode("utf-8"), headers=headers, method="POST"
        )
        with urllib.request.urlopen(req, timeout=15) as response:
            res_data = json.loads(response.read().decode("utf-8"))
            contrato_text = res_data["candidates"][0]["content"]["parts"][0]["text"]

            # Crear documento Word con encabezado
            doc = crear_documento_con_encabezado()
            doc.add_heading("CONTRATO DE LICENCIAMIENTO DE SERVICIOS PROFESIONALES", level=1)
            doc.add_paragraph()

            # Agregar el texto del contrato línea por línea
            for linea in contrato_text.split("\n"):
                linea = linea.strip()
                if linea:
                    doc.add_paragraph(linea)

            return guardar_contrato_docx(doc, orden_id, clon_id)

    except Exception as e:
        print(f"[GEMINI] Error generando contrato: {e}")
        return generar_contrato_default(
            orden_id,
            cliente_email,
            clon_id,
            clon_nombre,
            especialidad,
            cantidad_horas,
            monto_total,
            comision,
        )


def generar_contrato_default(
    orden_id,
    cliente_email,
    clon_id,
    clon_nombre,
    especialidad,
    cantidad_horas,
    monto_total,
    comision,
):
    """
    Contrato por defecto en formato Word si Gemini no está disponible.
    """

    fecha_inicio = datetime.now()
    fecha_fin = fecha_inicio + timedelta(days=30)

    # Crear documento con encabezado estándar
    doc = crear_documento_con_encabezado()

    # ── TÍTULO ──
    titulo = doc.add_heading(
        "CONTRATO DE LICENCIAMIENTO DE SERVICIOS PROFESIONALES", level=1
    )
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── DATOS DE LA ORDEN ──
    doc.add_paragraph()

    p_fecha = doc.add_paragraph()
    run_f = p_fecha.add_run(f"FECHA: {fecha_inicio.strftime('%d/%m/%Y')}")
    run_f.bold = True

    p_orden = doc.add_paragraph()
    run_o = p_orden.add_run(f"ID DE ORDEN: {orden_id}")
    run_o.bold = True

    doc.add_paragraph()

    # ── PARTES ──
    doc.add_heading("PARTES", level=2)

    p_lic = doc.add_paragraph()
    p_lic.add_run("1. LICENCIANTE: ").bold = True
    p_lic.add_run(clon_nombre)

    p_lic_esp = doc.add_paragraph()
    p_lic_esp.add_run(f"   - Especialidad: {especialidad}")
    p_lic_esp.add_run(f"\n   - Identificado como: {clon_id}")

    p_cli = doc.add_paragraph()
    p_cli.add_run("2. LICENCIATARIO (CLIENTE): ").bold = True
    p_cli.add_run(cliente_email)

    doc.add_paragraph()

    # ── ARTÍCULO 1 ──
    doc.add_heading("ARTÍCULO 1: OBJETO DEL CONTRATO", level=2)

    doc.add_paragraph(
        f"El Licenciante se compromete a proporcionar servicios profesionales de asesoría "
        f"y consultoría en el área de {especialidad} por un total de {cantidad_horas} horas "
        f"de trabajo."
    )

    doc.add_paragraph(
        "Los servicios se prestarán de conformidad con los más altos estándares de la "
        "industria y de acuerdo con la ley aplicable."
    )

    doc.add_paragraph()

    # ── ARTÍCULO 2 ──
    doc.add_heading("ARTÍCULO 2: TÉRMINOS ECONÓMICOS", level=2)

    tarifa_hora = monto_total / cantidad_horas if cantidad_horas > 0 else 0

    economicos = [
        f"- Tarifa por Hora: ${tarifa_hora:.2f} USD",
        f"- Horas Contratadas: {cantidad_horas} horas",
        f"- Monto Total: ${monto_total:.2f} USD",
        f"- Comisión de Plataforma: ${comision:.2f} USD (incluida en el total)",
        f"- Total a Pagar: ${monto_total:.2f} USD",
    ]

    for linea in economicos:
        doc.add_paragraph(linea)

    doc.add_paragraph("El pago debe realizarse según las indicaciones de SkillTwin.")

    doc.add_paragraph()

    # ── ARTÍCULO 3 ──
    doc.add_heading("ARTÍCULO 3: PLAZO", level=2)

    doc.add_paragraph(
        f"Este contrato es válido desde {fecha_inicio.strftime('%d/%m/%Y')} hasta "
        f"{fecha_fin.strftime('%d/%m/%Y')}."
    )

    doc.add_paragraph()

    # ── ARTÍCULO 4 ──
    doc.add_heading("ARTÍCULO 4: CONFIDENCIALIDAD", level=2)

    doc.add_paragraph(
        "Ambas partes se comprometen a mantener en confidencialidad toda la información "
        "compartida durante la prestación de servicios, excepto cuando lo autorice la ley "
        "o sea necesario para cumplir obligaciones legales."
    )

    doc.add_paragraph()

    # ── ARTÍCULO 5 ──
    doc.add_heading("ARTÍCULO 5: LIMITACIÓN DE RESPONSABILIDAD", level=2)

    doc.add_paragraph(
        "SkillTwin actúa como plataforma intermediaria. Ni SkillTwin ni el Licenciante "
        "serán responsables por daños indirectos, incidentales o consecuentes que puedan "
        "derivarse del uso de los servicios."
    )

    doc.add_paragraph()

    # ── ARTÍCULO 6 ──
    doc.add_heading("ARTÍCULO 6: TERMINACIÓN", level=2)

    doc.add_paragraph(
        "Este contrato puede ser terminado por cualquiera de las partes con notificación "
        "escrita con 48 horas de anticipación. En caso de terminación anticipada, se "
        "facturarán solo las horas trabajadas."
    )

    doc.add_paragraph()

    # ── ARTÍCULO 7 ──
    doc.add_heading("ARTÍCULO 7: LEY APLICABLE", level=2)

    doc.add_paragraph(
        "Este contrato se rige por las leyes del país donde SkillTwin está constituida y "
        "por los términos de servicio de la plataforma SkillTwin."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # ── FIRMAS ──
    p_firma_titulo = doc.add_paragraph()
    p_firma_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ft = p_firma_titulo.add_run("FIRMAS DIGITALES:")
    run_ft.bold = True
    run_ft.font.size = Pt(12)

    doc.add_paragraph()

    # Licenciante
    p_lic_firma = doc.add_paragraph()
    p_lic_firma.add_run(f"Licenciante: {clon_nombre}").bold = True
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("✓ Digitalmente Firmado")

    doc.add_paragraph()

    # Licenciatario
    p_cli_firma = doc.add_paragraph()
    p_cli_firma.add_run(f"Licenciatario: {cliente_email}").bold = True
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("✓ Digitalmente Firmado")

    doc.add_paragraph()

    # SkillTwin
    p_st_firma = doc.add_paragraph()
    p_st_firma.add_run("SkillTwin (Plataforma)").bold = True
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("✓ Validado por Sistema")

    # ── PIE ──
    doc.add_paragraph()
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = p_pie.add_run("SkillTwin © 2026 — Gemelos Digitales Profesionales")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(150, 150, 150)
    run_pie.italic = True

    # Guardar y retornar ruta
    return guardar_contrato_docx(doc, orden_id, clon_id)
