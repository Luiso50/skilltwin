from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

title = doc.add_heading('SkillTwin', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Gemelos Digitales Profesionales')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# --- MISION ---
doc.add_heading('Misión', level=1)
doc.add_paragraph(
    'Transformar el conocimiento experto de profesionales en gemelos digitales operables, '
    'brindando una plataforma integral que permita licenciar, monetizar y gestionar habilidades '
    'digitales con total trazabilidad, cumplimiento legal y visibilidad financiera.'
)

# --- VISION ---
doc.add_heading('Visión', level=1)
doc.add_paragraph(
    'Ser la infraestructura de referencia para el licenciamiento de talento digital con IA, '
    'convirtiendo la experiencia profesional en un activo repetible, escalable y gobernado, '
    'disponible para empresas, consultoras y profesionales que buscan ampliar su impacto '
    'mediante gemelos digitales inteligentes.'
)

# --- VALORES ---
doc.add_heading('Valores', level=1)
valores = [
    ('Innovación', 'Aplicamos IA de vanguardia para convertir conocimiento en activos digitales.'),
    ('Transparencia', 'Cada operación es auditable, desde contratos hasta flujos financieros.'),
    ('Cumplimiento Legal', 'Diseñamos nuestra plataforma con privacidad, ética y licenciamiento justo.'),
    ('Orientación al Cliente', 'El éxito de nuestros usuarios es el centro de nuestras decisiones.'),
]
for titulo, desc in valores:
    p = doc.add_paragraph()
    run_t = p.add_run(f'• {titulo}: ')
    run_t.bold = True
    p.add_run(desc)

# --- QUE ES SKILLTWIN ---
doc.add_heading('¿Qué es SkillTwin?', level=1)
doc.add_paragraph(
    'SkillTwin es una plataforma que convierte el conocimiento experto en gemelos digitales '
    'inteligentes. Cada gemelo digital (clone) representa a un profesional real y puede '
    'responder consultas, asesorar clientes y participar en flujos operativos automatizados, '
    'todo dentro de un ecosistema seguro y escalable.'
)

# --- QUE PUEDE HACER EL CLIENTE ---
doc.add_heading('¿Qué puede hacer un cliente con SkillTwin?', level=1)

servicios = [
    ('Consultar Gemelos Digitales Accede a clones de IA especializados en diversas industrias '
     '(COBOL, Finanzas, Ciberseguridad, UX, Data Science, Legal, Ventas) para obtener asesoría '
     'experta en tiempo real.'),
    ('Agendar una Demo Estratégica Solicita una sesión guiada para conocer el producto, revisar '
     'la arquitectura y validar cómo SkillTwin puede aplicarse a tu caso de uso específico.'),
    ('Lanzar un Piloto Corporativo Implementa un gemelo digital personalizado dentro de tu '
     'organización, con branding propio, portal de clientes y supervisión operativa.'),
    ('Obtener una Licencia de Concepto Evoluciona el prototipo hacia una solución interna o '
     'oferta comercial, con adaptación por sector e integración con tus operaciones existentes.'),
    ('Gestionar Contratos y Cumplimiento Accede a generación automatizada de acuerdos de '
     'licencia, con trazabilidad legal y estructura de privacidad por diseño.'),
    ('Monitorear el Rendimiento Financiero Visualiza dashboards de flujo de caja, cuentas por '
     'cobrar y pagar, y recibe alertas de pagos y cobranza para optimizar la operación.'),
    ('Realizar Pagos de Forma Segura Paga mediante transferencia bancaria, Bizum o tarjeta de '
     'crédito a través de Stripe, con facturación completa para empresas.'),
    ('Contactar al Equipo Soporte Utiliza los canales de WhatsApp, email o formulario web '
     'para solicitar demos, pilotos o asistencia técnica personalizada.'),
]

for servicio in servicios:
    doc.add_paragraph(servicio, style='List Bullet')

# --- CLONES DISPONIBLES ---
doc.add_heading('Gemelos Digitales Disponibles', level=1)

# Tabla de clones
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Nombre'
hdr_cells[2].text = 'Especialidad'

clones = [
    ('rsanchez_cobol', 'Roberto Sánchez', 'Programador Senior de COBOL'),
    ('ana_finanzas', 'Ana Gómez', 'Asesora de Finanzas Personales'),
    ('carlos_ciberseguridad', 'Carlos Mendoza', 'Experto en Ciberseguridad'),
    ('laura_ux', 'Laura Fernández', 'Diseñadora UX/UI'),
    ('pedro_data', 'Pedro Ruiz', 'Data Scientist'),
    ('maria_legal', 'María Torres', 'Abogada Tech / DPO'),
    ('diego_ventas', 'Diego Vargas', 'Director Comercial B2B'),
    ('fernando_telemedicina', 'Dr. Fernando López', 'Telemedicina y Salud Digital'),
    ('patricia_cloud', 'Patricia Morales', 'Arquitectura de Nube y DevOps'),
    ('alejandro_patentes', 'Alejandro Ríos', 'Propiedad Intelectual y Patentes'),
    ('valentina_rrhh', 'Valentina Herrera', 'Recursos Humanos y Talento Digital'),
    ('sebastian_manufactura', 'Sebastián Vargas', 'Manufactura y Supply Chain'),
]

for id_clone, nombre, especialidad in clones:
    row_cells = table.add_row().cells
    row_cells[0].text = id_clone
    row_cells[1].text = nombre
    row_cells[2].text = especialidad

doc.add_paragraph()

# --- PLANES ---
doc.add_heading('Planes y Precios', level=1)

planes = [
    ('Demo Estratégica — $149 USD', 
     'Sesión guiada de producto, revisión de arquitectura y recomendaciones de siguiente paso.'),
    ('Piloto Corporativo — $790 USD', 
     'Configuración de un caso de uso, personalización básica de branding y soporte en implementación inicial.'),
    ('Licencia de Concepto — A medida', 
     'Adaptación por sector, integración con operación existente y roadmap técnico y comercial.'),
]

for titulo, desc in planes:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{titulo}\n')
    run_t.bold = True
    p.add_run(desc)

# --- CONTACTO ---
doc.add_heading('Contacto', level=1)
doc.add_paragraph('WhatsApp: +34 600 000 000')
doc.add_paragraph('Email: teamskiltwinhq@zohomail.com')
doc.add_paragraph('Horario: Lunes a viernes, 09:00 – 18:00')
doc.add_paragraph('Web: https://luiso50.github.io/skilltwin/')

# --- PIE ---
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = footer.add_run('SkillTwin © 2026 — Proyecto experimental de gemelos digitales.')
run_f.font.size = Pt(9)
run_f.font.color.rgb = RGBColor(150, 150, 150)

output_path = os.path.join(os.path.dirname(__file__), 'SkillTwin_Mision_Vision.docx')
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
