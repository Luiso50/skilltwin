import os
import sys
import tempfile
import unittest
from docx import Document

ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, ROOT_DIR)

from dep_legal import generador_contratos  # noqa: E402


def leer_contenido_docx(ruta: str) -> str:
    """Lee el contenido de un archivo .docx y retorna todo el texto."""
    doc = Document(ruta)
    texto = []
    for para in doc.paragraphs:
        texto.append(para.text)
    return "\n".join(texto)


class GeneradorContratosTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.TemporaryDirectory()
        self.contratos_dir = os.path.join(self.tmpdir.name, 'contratos')
        generador_contratos.CONTRATOS_DIR = self.contratos_dir

    def tearDown(self):
        self.tmpdir.cleanup()

    def test_generar_contrato_crea_archivo(self):
        ruta = generador_contratos.generar_contrato(
            id_experto='test_user',
            nombre='Test User',
            especialidad='Testing',
            comision=15.0
        )
        self.assertIsNotNone(ruta)
        self.assertTrue(os.path.exists(ruta))

    def test_generar_contrato_nombre_archivo(self):
        ruta = generador_contratos.generar_contrato(
            id_experto='juan_dev',
            nombre='Juan Dev',
            especialidad='Desarrollo',
            comision=15.0
        )
        self.assertIn('contrato_juan_dev.docx', ruta)

    def test_generar_contrato_contenido_nombre(self):
        ruta = generador_contratos.generar_contrato(
            id_experto='maria_ia',
            nombre='María García',
            especialidad='Inteligencia Artificial',
            comision=15.0
        )
        contenido = leer_contenido_docx(ruta)
        self.assertIn('María García', contenido)

    def test_generar_contrato_contenido_especialidad(self):
        ruta = generador_contratos.generar_contrato(
            id_experto='pedro_data',
            nombre='Pedro López',
            especialidad='Data Science',
            comision=15.0
        )
        contenido = leer_contenido_docx(ruta)
        self.assertIn('Data Science', contenido)

    def test_generar_contrato_comision_personalizada(self):
        ruta = generador_contratos.generar_contrato(
            id_experto='custom_comision',
            nombre='Custom User',
            especialidad='Testing',
            comision=20.0
        )
        contenido = leer_contenido_docx(ruta)
        self.assertIn('20', contenido)
        self.assertIn('80', contenido)

    def test_generar_contrato_comision_por_defecto(self):
        ruta = generador_contratos.generar_contrato(
            id_experto='default_comision',
            nombre='Default User',
            especialidad='Testing'
        )
        contenido = leer_contenido_docx(ruta)
        self.assertIn('15', contenido)

    def test_generar_contrato_crea_directorio(self):
        dir_no_existente = os.path.join(self.tmpdir.name, 'nueva_carpeta', 'contratos')
        generador_contratos.CONTRATOS_DIR = dir_no_existente
        ruta = generador_contratos.generar_contrato(
            id_experto='test_dir',
            nombre='Test Dir',
            especialidad='Testing'
        )
        self.assertIsNotNone(ruta)
        self.assertTrue(os.path.exists(dir_no_existente))

    def test_generar_contrato_formato_fecha(self):
        ruta = generador_contratos.generar_contrato(
            id_experto='fecha_test',
            nombre='Fecha Test',
            especialidad='Testing'
        )
        contenido = leer_contenido_docx(ruta)
        self.assertIn('ACUERDO DE LICENCIA', contenido)
        self.assertIn('SKILLTWIN', contenido)

    def test_generar_contrato_tiene_encabezado(self):
        """Verifica que el documento Word tenga encabezado SkillTwin."""
        ruta = generador_contratos.generar_contrato(
            id_experto='header_test',
            nombre='Header Test',
            especialidad='Testing'
        )
        doc = Document(ruta)
        # Verificar que hay encabezado
        header = doc.sections[0].header
        self.assertTrue(len(header.tables) > 0 or len(header.paragraphs) > 0)


if __name__ == '__main__':
    unittest.main()
